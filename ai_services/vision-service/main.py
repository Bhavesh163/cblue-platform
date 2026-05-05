"""
CBLUE AI Vision Service — Portfolio & KYC Document Digestion

Handles OCR and text extraction from uploaded files:
- JPEG/PNG: Tesseract OCR (tha+eng)
- PDF:      pdf2image → Tesseract OCR (tha+eng), fallback message if unreadable
- DOCX:     python-docx paragraph + table extraction
- XLSX:     openpyxl per-sheet extraction

Returns extracted text + verification score for AI evaluation.
"""

import io
import os
import tempfile
import uuid
from pathlib import Path
from datetime import datetime, timezone

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

app = FastAPI(title="CBLUE Vision Service", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── Storage for extracted texts ──
STORAGE_DIR = Path(os.getenv("VISION_STORAGE_DIR", "/tmp/cblue-vision"))
STORAGE_DIR.mkdir(parents=True, exist_ok=True)

# ── Supported MIME types ──
ALLOWED_TYPES = {
    "image/jpeg", "image/jpg", "image/png", "image/webp",
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # docx
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",  # xlsx
    "application/vnd.ms-excel",  # xls
    "text/plain",  # txt
}

MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB


class ExtractionResult(BaseModel):
    file_id: str
    filename: str
    file_type: str
    raw_text: str
    text_length: int
    extraction_method: str
    has_content: bool
    verification_hints: list[str]
    timestamp: str


class BatchResult(BaseModel):
    results: list[ExtractionResult]
    total_files: int
    total_text_length: int
    content_score: float  # 0-100 based on extractable content quality


# ── OCR via Tesseract ──

def ocr_image(image_bytes: bytes, lang: str = "tha+eng") -> str:
    """Extract text from image bytes using Typhoon AI Vision."""
    try:
        from typhoon import ocr_image_typhoon
        return ocr_image_typhoon(image_bytes)
    except Exception as e:
        return f"[OCR error: {e}]"


def ocr_pdf(pdf_bytes: bytes, lang: str = "tha+eng") -> str:
    """Extract text from PDF pages using pdf2image + Typhoon AI."""
    try:
        from pdf2image import convert_from_bytes
        from typhoon import ocr_image_typhoon
        import io

        # Limit to 3 pages to avoid excessive AI API costs
        pages = convert_from_bytes(pdf_bytes, dpi=150, first_page=1, last_page=3)
        texts = []
        for i, page_img in enumerate(pages):
            if page_img.mode != "RGB":
                page_img = page_img.convert("RGB")
            buf = io.BytesIO()
            page_img.save(buf, format='JPEG')
            page_text = ocr_image_typhoon(buf.getvalue())
            if page_text.strip():
                texts.append(f"__PAGE_{i + 1}__\n{page_text.strip()}")
        return "\n\n".join(texts) if texts else ""
    except ImportError:
        return "[PDF OCR unavailable — pdf2image/poppler not installed]"
    except Exception as e:
        return f"[PDF OCR error: {e}]"


# ── DOCX extraction ──

def extract_docx(docx_bytes: bytes) -> str:
    """Extract text from DOCX including paragraphs and tables."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(docx_bytes))
        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for i, table in enumerate(doc.tables):
            parts.append(f"__TABLE_START_{i + 1}__")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append("\t".join(cells))
            parts.append(f"__TABLE_END_{i + 1}__")

        return "\n".join(parts)
    except ImportError:
        return "[DOCX extraction unavailable — python-docx not installed]"
    except Exception as e:
        return f"[DOCX error: {e}]"


# ── Excel extraction ──

def extract_excel(xlsx_bytes: bytes) -> str:
    """Extract text from Excel workbook, per-sheet."""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(xlsx_bytes), read_only=True, data_only=True)
        parts: list[str] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"__SHEET:{sheet_name}__")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    parts.append("\t".join(cells))

        wb.close()
        return "\n".join(parts)
    except ImportError:
        return "[Excel extraction unavailable — openpyxl not installed]"
    except Exception as e:
        return f"[Excel error: {e}]"


# ── TXT extraction ──

def extract_txt(txt_bytes: bytes) -> str:
    """Extract text from plain text file, fallback encoding."""
    try:
        return txt_bytes.decode('utf-8')
    except UnicodeDecodeError:
        try:
            return txt_bytes.decode('cp874')
        except Exception as e:
            return f"[TXT decode error: {e}]"


# ── Verification hint extraction ──

def extract_hints(text: str, filename: str) -> list[str]:
    """Analyze extracted text for verification-relevant hints."""
    hints: list[str] = []
    lower = text.lower()

    # Check for professional credentials
    credential_keywords = [
        ("license", "Professional license reference found"),
        ("certificate", "Certificate reference found"),
        ("ใบอนุญาต", "Thai professional license reference"),
        ("ใบรับรอง", "Thai certificate reference"),
        ("วุฒิบัตร", "Thai qualification certificate"),
        ("ทะเบียน", "Thai registration reference"),
        ("สภา", "Thai professional council reference"),
        ("วิศวกรรม", "Engineering reference"),
        ("สถาปัตยกรรม", "Architecture reference"),
        ("certified", "Certification claim found"),
        ("registered", "Registration claim found"),
        ("project", "Project reference found"),
        ("experience", "Experience claim found"),
        ("ประสบการณ์", "Thai experience reference"),
        ("โครงการ", "Thai project reference"),
    ]

    for keyword, hint in credential_keywords:
        if keyword in lower:
            hints.append(hint)

    # Check for company/org references
    company_keywords = ["co.,ltd", "company", "บริษัท", "ห้างหุ้นส่วน", "corporation"]
    for kw in company_keywords:
        if kw in lower:
            hints.append(f"Company/organization reference: '{kw}'")
            break

    # Check for date patterns (suggests real documents)
    import re
    date_patterns = re.findall(r'\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4}', text)
    if date_patterns:
        hints.append(f"Date references found ({len(date_patterns)} instances)")

    # Check for numeric patterns (prices, IDs, phone numbers)
    numbers = re.findall(r'\d{3,}', text)
    if len(numbers) > 3:
        hints.append("Multiple numeric values found (possible price list or ID references)")

    if not hints:
        hints.append("No specific credential references detected in extracted text")

    return hints


def save_extraction(file_id: str, raw_text: str, filename: str) -> None:
    """Save extracted text for audit trail."""
    file_dir = STORAGE_DIR / file_id
    file_dir.mkdir(parents=True, exist_ok=True)

    (file_dir / "raw_ocr.txt").write_text(raw_text, encoding="utf-8")
    (file_dir / "raw_text.txt").write_text(raw_text, encoding="utf-8")
    (file_dir / "metadata.txt").write_text(
        f"filename: {filename}\nextracted: {datetime.now(timezone.utc).isoformat()}\ntext_length: {len(raw_text)}\n",
        encoding="utf-8",
    )


# ── API Endpoints ──

@app.post("/extract", response_model=ExtractionResult)
async def extract_single(file: UploadFile = File(...)):
    """Extract text from a single uploaded file."""
    if file.content_type and file.content_type not in ALLOWED_TYPES:
        raise HTTPException(400, f"Unsupported file type: {file.content_type}")

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(413, "File too large (max 20MB)")

    file_id = str(uuid.uuid4())
    filename = file.filename or "unknown"
    ct = file.content_type or ""
    raw_text = ""
    method = "unknown"

    if ct.startswith("image/"):
        raw_text = ocr_image(content)
        method = "tesseract_ocr"
    elif ct == "application/pdf":
        raw_text = ocr_pdf(content)
        method = "pdf2image_tesseract"
        if not raw_text or raw_text.startswith("["):
            raw_text = raw_text or "Sorry, please fill in the input"
    elif "wordprocessingml" in ct:
        raw_text = extract_docx(content)
        method = "python_docx"
    elif "spreadsheetml" in ct or "ms-excel" in ct:
        raw_text = extract_excel(content)
        method = "openpyxl"
    elif ct == "text/plain":
        raw_text = extract_txt(content)
        method = "txt_decode"
    else:
        # Fallback: try OCR on anything
        raw_text = ocr_image(content)
        method = "fallback_ocr"

    has_content = bool(raw_text) and not raw_text.startswith("[")
    hints = extract_hints(raw_text, filename) if has_content else ["No extractable content"]

    save_extraction(file_id, raw_text, filename)

    return ExtractionResult(
        file_id=file_id,
        filename=filename,
        file_type=ct,
        raw_text=raw_text,
        text_length=len(raw_text),
        extraction_method=method,
        has_content=has_content,
        verification_hints=hints,
        timestamp=datetime.now(timezone.utc).isoformat(),
    )


@app.post("/extract-batch", response_model=BatchResult)
async def extract_batch(files: list[UploadFile] = File(...)):
    """Extract text from multiple uploaded files (portfolio batch)."""
    if len(files) > 10:
        raise HTTPException(400, "Maximum 10 files per batch")

    results: list[ExtractionResult] = []
    for f in files:
        if f.content_type and f.content_type not in ALLOWED_TYPES:
            results.append(ExtractionResult(
                file_id=str(uuid.uuid4()),
                filename=f.filename or "unknown",
                file_type=f.content_type or "",
                raw_text=f"[Unsupported type: {f.content_type}]",
                text_length=0,
                extraction_method="skipped",
                has_content=False,
                verification_hints=["Unsupported file type"],
                timestamp=datetime.now(timezone.utc).isoformat(),
            ))
            continue

        content = await f.read()
        if len(content) > MAX_FILE_SIZE:
            results.append(ExtractionResult(
                file_id=str(uuid.uuid4()),
                filename=f.filename or "unknown",
                file_type=f.content_type or "",
                raw_text="[File too large]",
                text_length=0,
                extraction_method="skipped",
                has_content=False,
                verification_hints=["File exceeds 20MB limit"],
                timestamp=datetime.now(timezone.utc).isoformat(),
            ))
            continue

        file_id = str(uuid.uuid4())
        filename = f.filename or "unknown"
        ct = f.content_type or ""
        raw_text = ""
        method = "unknown"

        if ct.startswith("image/"):
            raw_text = ocr_image(content)
            method = "tesseract_ocr"
        elif ct == "application/pdf":
            raw_text = ocr_pdf(content)
            method = "pdf2image_tesseract"
            if not raw_text or raw_text.startswith("["):
                raw_text = raw_text or "Sorry, please fill in the input"
        elif "wordprocessingml" in ct:
            raw_text = extract_docx(content)
            method = "python_docx"
        elif "spreadsheetml" in ct or "ms-excel" in ct:
            raw_text = extract_excel(content)
            method = "openpyxl"
        elif ct == "text/plain":
            raw_text = extract_txt(content)
            method = "txt_decode"
        else:
            raw_text = ocr_image(content)
            method = "fallback_ocr"

        has_content = bool(raw_text) and not raw_text.startswith("[")
        hints = extract_hints(raw_text, filename) if has_content else ["No extractable content"]
        save_extraction(file_id, raw_text, filename)

        results.append(ExtractionResult(
            file_id=file_id,
            filename=filename,
            file_type=ct,
            raw_text=raw_text,
            text_length=len(raw_text),
            extraction_method=method,
            has_content=has_content,
            verification_hints=hints,
            timestamp=datetime.now(timezone.utc).isoformat(),
        ))

    total_text = sum(r.text_length for r in results)
    files_with_content = sum(1 for r in results if r.has_content)
    content_score = (files_with_content / max(len(results), 1)) * 100 if results else 0

    return BatchResult(
        results=results,
        total_files=len(results),
        total_text_length=total_text,
        content_score=round(content_score, 1),
    )



class RegistrationData(BaseModel):
    experience_years: int
    skills: list[str]
    profile_completeness: int
    has_price_list: bool
    ocr_texts: list[str]

@app.post("/score-registration")
async def score_registration(data: RegistrationData):
    from typhoon import ocr_image_typhoon
    import httpx
    # Use LLM directly to score
    api_key = "sk-9QMPk7A1MYzAGi7dgjTldnK7DB9Ion8fRkiGC1uIX632Fg9y"
    llm_headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }

    combined_ocr = "
".join(data.ocr_texts)[:5000]

    system_prompt = """You are a world-class enterprise AI Verification Assistant for CBLUE platform.
Evaluate this partner registration and return a STRICT JSON output with scores.
Score breakdown (100 max):
- Experience (0-25)
- Skills Breadth (0-15)
- KYC Verification (0-15)
- Portfolio & Evidence (0-15)
- Profile Completeness (0-10)
- Price List (0-10)
- Credential Verification (0-10)
Also provide tier: Economy, Standard, Corporate, Specialist, or Expert.
Output JSON MUST look exactly like:
{ "score": 85, "breakdown": { "experience": 20, "skills": 12, "kyc": 15, "portfolio": 10, "profile": 10, "price": 10, "credential": 8 }, "tier": "Specialist", "hints": ["Hint 1"] }
"""
    user_prompt = f"""
Partner stats:
Exp: {data.experience_years}
Skills count: {len(data.skills)}
Completeness: {data.profile_completeness}/100
Price list: {data.has_price_list}
OCR Texts (Portfolio/KYC):
{combined_ocr}
    """
    llm_payload = {
        "model": "Typhoon-v2.5-30b-a3b-instruct",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        "max_tokens": 500,
        "temperature": 0.1
    }
    try:
        with httpx.Client(timeout=60.0) as client:
            llm_res = client.post("https://api.opentyphoon.ai/v1/chat/completions", headers=llm_headers, json=llm_payload)
            llm_res.raise_for_status()
            llm_data = llm_res.json()
            import json
            raw = llm_data['choices'][0]['message']['content'].strip()
            # find JSON block
            import re
            m = re.search(r'\{.*\}', raw, re.DOTALL)
            if m: raw = m.group(0)
            return json.loads(raw)
    except Exception as e:
        return {
            "score": 60,
            "breakdown": {"experience": 15, "skills": 10, "kyc": 10, "portfolio": 5, "profile": 10, "price": 10, "credential": 0},
            "tier": "Standard",
            "hints": [f"AI scoring failed gracefully: {e}"]
        }


@app.get("/health")
async def health():
    """Health check endpoint."""
    tesseract_ok = False
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        tesseract_ok = True
    except Exception:
        pass

    return {
        "status": "ok",
        "service": "vision-service",
        "tesseract_available": tesseract_ok,
        "storage_dir": str(STORAGE_DIR),
    }
